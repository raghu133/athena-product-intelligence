"""Generate a few realistic PDF and DOCX source documents into data/raw.

These complement the JSON corpus with real-world file formats:
  * research-market-2026.pdf     — a market/research document
  * competitor-deepdive.pdf      — a competitor analysis deck-style doc
  * prd-global-search.docx       — a PRD exported from Word

They reference the SAME customers, features, and product areas as the JSON
corpus (SSO, global search, Northwind, Acme, etc.) so retrieval connects across
formats. They arrive WITHOUT structured metadata, exercising the LLM enrichment
path during ingestion.

Run: python scripts/make_sample_docs.py
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from athena.core.config import RAW_DIR


def _write_pdf(path: Path, title: str, paragraphs: list[str]) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    for p in paragraphs:
        if p.startswith("## "):
            story.append(Paragraph(p[3:], styles["Heading2"]))
        else:
            story.append(Paragraph(p, styles["BodyText"]))
        story.append(Spacer(1, 8))
    SimpleDocTemplate(str(path), pagesize=A4).build(story)


def _write_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    from docx import Document as Docx
    doc = Docx()
    doc.add_heading(title, level=0)
    for p in paragraphs:
        if p.startswith("## "):
            doc.add_heading(p[3:], level=2)
        else:
            doc.add_paragraph(p)
    doc.save(str(path))


def main() -> int:
    RAW_DIR.mkdir(parents=True, exist_ok=True)

    _write_pdf(
        RAW_DIR / "research-market-2026.pdf",
        "Market Research: Collaboration & Analytics SaaS (2026)",
        [
            "## Executive Summary",
            "Enterprise buyers increasingly gate purchases on security and identity "
            "features. SSO/SAML and SCIM provisioning are now table-stakes for "
            "deals above 500 seats. Vendors lacking these lose evaluations late in "
            "the cycle, after significant sales investment.",
            "## Key Findings",
            "1. Cross-workspace search is a top-3 requested capability among "
            "mid-market and enterprise teams; fragmented search is a leading cause "
            "of daily friction.",
            "2. Notification overload drives disengagement. Buyers want granular, "
            "per-project controls rather than all-or-nothing email.",
            "3. Reliability of integrations (Slack, webhooks) materially affects "
            "renewal sentiment; silent webhook failures are a recurring complaint.",
            "## Recommendation",
            "Prioritize SSO/SAML, SCIM, and global search to defend enterprise "
            "deals against competitors such as Tandem and GridWork, which already "
            "ship these out of the box.",
        ],
    )

    _write_pdf(
        RAW_DIR / "competitor-deepdive.pdf",
        "Competitor Deep Dive: Tandem vs GridWork vs Flowdesk",
        [
            "## Overview",
            "This deep dive compares Flowdesk against Tandem and GridWork across "
            "the capabilities that most influence enterprise evaluations.",
            "## Identity & Access",
            "Both Tandem and GridWork ship SSO/SAML and SCIM. Flowdesk currently "
            "does not, which blocks company-wide rollouts for customers like "
            "Northwind Trading and Cobalt Health.",
            "## Search",
            "GridWork offers global search across workspaces. Flowdesk search is "
            "scoped to a single workspace, a frequent source of customer frustration.",
            "## Threat Assessment",
            "We are most exposed in enterprise deals requiring SSO and "
            "cross-workspace search. Closing these gaps is the highest-leverage "
            "competitive move.",
        ],
    )

    _write_docx(
        RAW_DIR / "prd-global-search.docx",
        "PRD: Global Search Across Workspaces",
        [
            "## Status",
            "Draft — proposed for next-quarter prioritization.",
            "## Problem",
            "Customers repeatedly request search that spans all their workspaces. "
            "Today search is limited to the current workspace, so users cannot "
            "find documents owned by other teams. This appears in support tickets, "
            "NPS feedback, and customer interviews (e.g. Acme Robotics, Quill "
            "Publishing).",
            "## Proposed Solution",
            "Introduce a global search index across a customer's workspaces with "
            "permission-aware result filtering.",
            "## Success Metrics",
            "Reduce Search-related complaints by 40%; achieve >50% adoption of "
            "global search within 60 days of launch.",
            "## Competitive Note",
            "GridWork already offers this; shipping it defends at-risk enterprise "
            "renewals.",
        ],
    )

    made = ["research-market-2026.pdf", "competitor-deepdive.pdf", "prd-global-search.docx"]
    print(f"[ok] wrote {len(made)} mixed-format docs to {RAW_DIR}:")
    for m in made:
        print("   ", m)
    return 0


if __name__ == "__main__":
    sys.exit(main())
