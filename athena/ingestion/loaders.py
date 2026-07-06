"""Load raw source documents from data/raw into typed Document objects.

Real organizational knowledge arrives in mixed formats, so the loader is
format-agnostic. It handles:
  * .json  — the structured synthetic corpus (rich metadata)
  * .md / .txt — plain text drop-ins
  * .pdf   — e.g. research documents, competitor decks (parsed with pypdf)
  * .docx  — e.g. PRDs, meeting notes exported from Word (python-docx)

PDF/DOCX/MD files carry no structured metadata, which is exactly why the
ingestion pipeline enriches them with an LLM extractor (see enrich.py) — so this
also demonstrates the unstructured-input path, not just the clean JSON path.
"""
from __future__ import annotations

import json
from pathlib import Path

from athena.core import config
from athena.core.schemas import Document


def _infer_source_type(name: str) -> str:
    stem = name.split("-")[0].lower()
    mapping = {
        "ticket": "support_ticket", "feedback": "customer_feedback",
        "prd": "prd", "meeting": "meeting_notes", "gh": "github_issue",
        "release": "release_notes", "interview": "customer_interview",
        "competitor": "competitor_analysis", "research": "research_document",
    }
    return mapping.get(stem, "document")


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    return "\n\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _read_docx(path: Path) -> str:
    from docx import Document as Docx
    doc = Docx(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()


def load_documents(raw_dir: Path | None = None) -> list[Document]:
    raw_dir = raw_dir or config.RAW_DIR
    docs: list[Document] = []
    for path in sorted(raw_dir.glob("*")):
        suffix = path.suffix.lower()
        if suffix == ".json":
            data = json.loads(path.read_text(encoding="utf-8"))
            # A file may hold a single document object OR an array of them
            # (source-type exports like customer_feedback.json).
            items = data if isinstance(data, list) else [data]
            for item in items:
                docs.append(Document(
                    doc_id=item["doc_id"],
                    source_type=item["source_type"],
                    title=item.get("title", item["doc_id"]),
                    text=item["text"],
                    created_at=item.get("created_at", "unknown"),
                    metadata=item.get("metadata", {}),
                ))
        elif suffix in (".md", ".txt", ".pdf", ".docx"):
            try:
                if suffix == ".pdf":
                    text = _read_pdf(path)
                elif suffix == ".docx":
                    text = _read_docx(path)
                else:
                    text = path.read_text(encoding="utf-8")
            except Exception:  # noqa: BLE001 — skip unreadable files, never crash ingest
                continue
            if not text:
                continue
            docs.append(Document(
                doc_id=path.stem,
                source_type=_infer_source_type(path.stem),
                title=path.stem.replace("_", " ").replace("-", " ").title(),
                text=text,
                created_at="unknown",
                metadata={"filename": path.name, "format": suffix.lstrip(".")},
            ))
    return docs
